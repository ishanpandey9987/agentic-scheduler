"""
Parsing Agent - Extracts schedule from PDF, DOCX, Images using Vision AI
Supports multiple document formats with Azure OpenAI GPT-4 Vision
"""
import json
import base64
import os
import sys
import requests
from typing import List, Optional
from pathlib import Path

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

# Optional imports for document processing
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from pdf2image import convert_from_path
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

from models.schedule_item import ScheduleItem, EventType
from config.settings import (
    AZURE_OPENAI_API_KEY,
    AZURE_OPENAI_ENDPOINT,
    OPENAI_DEPLOYMENT_NAME,
    OPENAI_API_VERSION
)


class ParsingAgent:
    """
    Extracts schedule information from various document formats:
    - PDF (text extraction + vision for scanned PDFs)
    - DOCX (Word documents)
    - Images (PNG, JPG, JPEG)
    - Plain text
    
    Uses Azure OpenAI GPT-4 Vision for intelligent extraction.
    """
    
    def __init__(self):
        self.api_url = f"{AZURE_OPENAI_ENDPOINT.rstrip('/')}/openai/deployments/{OPENAI_DEPLOYMENT_NAME}/chat/completions?api-version={OPENAI_API_VERSION}"
        self.headers = {
            "Content-Type": "application/json",
            "api-key": AZURE_OPENAI_API_KEY
        }
        self.system_prompt = self._get_system_prompt()
    
    def _get_system_prompt(self) -> str:
        """Get the system prompt for schedule extraction"""
        return """You are a schedule parsing assistant. Extract ALL events/classes from the provided content.

IMPORTANT: Return ONLY a valid JSON array, no other text or markdown.

Each event must have this exact structure:
[
    {
        "course": "Course/Event Name",
        "type": "lecture|lab|exam|meeting|practice|other",
        "location": "Room/Location (use empty string if not found)",
        "date": "YYYY-MM-DD (use current week dates if only day names given)",
        "from": "HH:MM (24-hour format)",
        "to": "HH:MM (24-hour format)"
    }
]

Guidelines:
- Extract EVERY event you can find
- If day names (Monday, Tuesday) are used without dates, use dates from the current week starting from today
- Convert 12-hour times (2:30 PM) to 24-hour format (14:30)
- If location is missing, use empty string ""
- If event type is unclear, use "other"
- Today's date is provided in the user message for reference

Return ONLY the JSON array, nothing else."""

    def parse_document(self, file_path: str) -> List[ScheduleItem]:
        """
        Main method to parse any supported document type.
        Automatically detects file type and uses appropriate method.
        
        Args:
            file_path: Path to the document (PDF, DOCX, PNG, JPG, etc.)
            
        Returns:
            List of ScheduleItem objects
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_schedule_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_schedule_from_docx(file_path)
        elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            return self.extract_schedule_from_image(file_path)
        elif extension in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return self.parse_schedule_text(f.read())
        else:
            print(f"❌ Unsupported file format: {extension}")
            return []

    def extract_schedule_from_pdf(self, pdf_path: str) -> List[ScheduleItem]:
        """
        Extracts schedule information from a PDF file.
        Uses text extraction first, falls back to vision for scanned PDFs.
        
        Args:
            pdf_path: The path to the PDF file.

        Returns:
            List of extracted ScheduleItem objects.
        """
        print(f"📄 Processing PDF: {pdf_path}")
        
        # Try text extraction first (faster for text-based PDFs)
        text_content = self._extract_pdf_text(pdf_path)
        
        if text_content and len(text_content.strip()) > 50:
            print("   Using text extraction method...")
            return self.parse_schedule_text(text_content)
        
        # Fall back to vision for scanned PDFs or image-based PDFs
        print("   PDF appears to be scanned/image-based, using vision...")
        return self._extract_pdf_with_vision(pdf_path)
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        """Extract text from PDF using PyPDF2"""
        if not HAS_PYPDF2:
            print("   ⚠️ PyPDF2 not installed, skipping text extraction")
            return ""
        
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            print(f"   ⚠️ PDF text extraction failed: {e}")
            return ""
    
    def _extract_pdf_with_vision(self, pdf_path: str) -> List[ScheduleItem]:
        """Convert PDF to images and use vision API"""
        if not HAS_PDF2IMAGE:
            print("   ❌ pdf2image not installed. Install with: pip install pdf2image")
            print("   Also install poppler: brew install poppler (macOS) or apt-get install poppler-utils (Linux)")
            return []
        
        try:
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=150)
            all_events = []
            
            for i, image in enumerate(images):
                print(f"   Processing page {i + 1}/{len(images)}...")
                
                # Save image temporarily
                temp_path = f"/tmp/pdf_page_{i}.png"
                image.save(temp_path, 'PNG')
                
                # Extract schedule from image
                events = self.extract_schedule_from_image(temp_path)
                all_events.extend(events)
                
                # Clean up
                os.remove(temp_path)
            
            # Remove duplicates based on course + date + time
            unique_events = self._deduplicate_events(all_events)
            return unique_events
            
        except Exception as e:
            print(f"   ❌ PDF vision extraction failed: {e}")
            return []

    def extract_schedule_from_image(self, image_path: str) -> List[ScheduleItem]:
        """
        Extracts schedule information from an image file using GPT-4 Vision.
        
        Args:
            image_path: The path to the image file.

        Returns:
            List of extracted ScheduleItem objects.
        """
        print(f"🖼️  Processing image: {image_path}")
        
        # Read and encode image
        try:
            with open(image_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")
        except Exception as e:
            print(f"   ❌ Failed to read image: {e}")
            return []
        
        # Determine media type
        extension = Path(image_path).suffix.lower()
        media_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.webp': 'image/webp'
        }
        media_type = media_types.get(extension, 'image/png')
        
        # Get current date for context
        from datetime import datetime
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        payload = {
            "messages": [
                {"role": "system", "content": self.system_prompt},
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"Today's date is {current_date}. Extract ALL schedule events from this image:"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{image_data}",
                                "detail": "high"  # Use high detail for better text recognition
                            }
                        }
                    ]
                }
            ],
            "temperature": 0.1,
            "max_tokens": 4000
        }
        
        return self._call_api_and_parse(payload)
    
    def extract_schedule_from_docx(self, docx_path: str) -> List[ScheduleItem]:
        """
        Extracts schedule information from a Word document.
        
        Args:
            docx_path: The path to the DOCX file.

        Returns:
            List of extracted ScheduleItem objects.
        """
        print(f"📝 Processing DOCX: {docx_path}")
        
        if not HAS_DOCX:
            print("   ❌ python-docx not installed. Install with: pip install python-docx")
            return []
        
        try:
            doc = DocxDocument(docx_path)
            
            # Extract text from paragraphs
            text_content = ""
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            
            # Extract text from tables (schedules are often in tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            if text_content.strip():
                return self.parse_schedule_text(text_content)
            else:
                print("   ⚠️ No text content found in DOCX")
                return []
                
        except Exception as e:
            print(f"   ❌ DOCX extraction failed: {e}")
            return []

    def parse_schedule_text(self, text: str) -> List[ScheduleItem]:
        """
        Parses schedule information from a text string using LLM.
        
        Args:
            text: The text containing schedule information.

        Returns:
            List of extracted ScheduleItem objects.
        """
        print("📝 Parsing text content...")
        
        if not text or len(text.strip()) < 10:
            print("   ⚠️ Text content too short")
            return []
        
        # Get current date for context
        from datetime import datetime
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        payload = {
            "messages": [
                {"role": "system", "content": self.system_prompt},
                {
                    "role": "user",
                    "content": f"Today's date is {current_date}. Extract ALL schedule events from this text:\n\n{text}"
                }
            ],
            "temperature": 0.1,
            "max_tokens": 4000
        }
        
        return self._call_api_and_parse(payload)
    
    def _call_api_and_parse(self, payload: dict) -> List[ScheduleItem]:
        """Call the API and parse the response into ScheduleItem objects"""
        try:
            response = requests.post(
                self.api_url, 
                headers=self.headers, 
                json=payload, 
                timeout=120
            )
            
            if response.status_code != 200:
                print(f"   ❌ API error {response.status_code}: {response.text[:200]}")
                return []
            
            result = response.json()
            content = result["choices"][0]["message"]["content"]
            
            # Clean up response (remove markdown code blocks if present)
            content = self._clean_json_response(content)
            
            # Parse JSON
            events_data = json.loads(content)
            
            # Convert to ScheduleItem objects
            events = []
            for e in events_data:
                try:
                    item = ScheduleItem.from_dict(e)
                    events.append(item)
                except Exception as ex:
                    print(f"   ⚠️ Skipping invalid event: {ex}")
            
            print(f"   ✅ Extracted {len(events)} events")
            return events
            
        except json.JSONDecodeError as e:
            print(f"   ❌ JSON parse error: {e}")
            print(f"   Response content: {content[:500] if 'content' in dir() else 'N/A'}")
            return []
        except Exception as e:
            print(f"   ❌ API call failed: {e}")
            return []
    
    def _clean_json_response(self, content: str) -> str:
        """Clean up JSON response from LLM"""
        content = content.strip()
        
        # Remove markdown code blocks
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]
        
        return content.strip()
    
    def _deduplicate_events(self, events: List[ScheduleItem]) -> List[ScheduleItem]:
        """Remove duplicate events based on key fields"""
        seen = set()
        unique = []
        
        for event in events:
            key = (event.course.lower(), event.date, event.start_time, event.end_time)
            if key not in seen:
                seen.add(key)
                unique.append(event)
        
        return unique
    
    def parse_schedule_url(self, url: str) -> List[ScheduleItem]:
        """
        Parse schedule from an image URL.
        
        Args:
            url: URL of the image containing the schedule
            
        Returns:
            List of extracted ScheduleItem objects
        """
        print(f"🌐 Processing URL: {url}")
        
        from datetime import datetime
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        payload = {
            "messages": [
                {"role": "system", "content": self.system_prompt},
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"Today's date is {current_date}. Extract ALL schedule events from this image:"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": url,
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            "temperature": 0.1,
            "max_tokens": 4000
        }
        
        return self._call_api_and_parse(payload)